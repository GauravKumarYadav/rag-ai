"""DOCX document processor.

Uses python-docx for Microsoft Word document extraction.
"""

from typing import Optional
from app.processors.base import DocumentProcessor, ProcessorConfig
from app.processors.registry import ProcessorRegistry


@ProcessorRegistry.register
class DOCXProcessor(DocumentProcessor):
    """Processor for Microsoft Word documents.
    
    Extracts text from .docx files, preserving paragraph structure.
    Also handles .doc files if possible (requires additional libraries).
    """
    
    name = "DOCX Processor"
    supported_extensions = [".docx", ".doc"]
    supported_mimetypes = [
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    ]
    
    def __init__(self, config: Optional[ProcessorConfig] = None):
        """Initialize DOCX processor and check dependencies."""
        super().__init__(config)
        self._docx_available = self._check_docx()
    
    def _check_docx(self) -> bool:
        """Check if python-docx is available."""
        try:
            import docx
            return True
        except ImportError:
            return False
    
    def extract(self, content: bytes, filename: str) -> str:
        """Extract text from Word documents.
        
        Args:
            content: Raw bytes of the document
            filename: Original filename
            
        Returns:
            Extracted text content
            
        Raises:
            ProcessorError: If extraction fails or library not installed
        """
        if not self._docx_available:
            raise ValueError(
                f"Cannot process '{filename}'. "
                "Please install python-docx: pip install python-docx"
            )
        
        import io
        import docx
        
        try:
            doc = docx.Document(io.BytesIO(content))
            
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        paragraphs.append(" | ".join(row_text))
            
            return "\n\n".join(paragraphs)
            
        except Exception as e:
            raise ValueError(
                f"Failed to extract text from '{filename}': {str(e)}"
            )
